from youtube_transcript_api import YouTubeTranscriptApi
import sys
from collections import Counter
import re
# from sklearn.feature_extraction.text import TfidfVectorizer
import spacy
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.probability import FreqDist
from docx import Document
from docx.shared import Inches
import os
import math

# id do video
id = sys.argv[1]
# title in docx
title = sys.argv[2]


# Carregar o modelo de idioma em inglês
nlp = spacy.load('en_core_web_sm')

srt = YouTubeTranscriptApi.get_transcript(str(id), languages=['en'])

texto = ''

for text in srt:
    texto = texto + text['text'] 


def remove_special_characters(text):
    # Substituir \xa0\n e \xa0\xa0 por espaços
    cleaned_text = re.sub(r'[\xa0\n]+', ' ', text)
    return cleaned_text


#removendo caracteris especiais
texto = remove_special_characters(texto)

# Pré-processamento do texto
stop_words = set(stopwords.words('english'))
lemmatizer = WordNetLemmatizer() 
doc = nlp(texto)


# Inicializar listas para armazenar verbos, adjetivos e pronomes
verbs = [lemmatizer.lemmatize(token.text) for token in doc if token.pos_ == 'VERB']
adjectives = [lemmatizer.lemmatize(token.text) for token in doc if token.pos_ == 'ADJ']
pronouns = [lemmatizer.lemmatize(token.text) for token in doc if token.pos_ == 'PRON']
noun = [lemmatizer.lemmatize(token.text) for token in doc if token.pos_ == 'NOUN']
conjuction = [lemmatizer.lemmatize(token.text) for token in doc if token.pos_ == 'CONJ']



def separa(dados, n):
    table = []
    base = []
    
    for i, palavra in enumerate(dados):
        base.append(palavra)
        if (i+1) % n == 0:
            table.append(base)
            base = []

    return table


def salva_palavras(palavras, file):
    n = math.floor(math.sqrt(len(set(palavras))))

    if n > 4:
        n = 4

    data = separa(set(palavras), n)

    if os.path.isfile(file):
        doc = Document(file)
    else:
        doc = Document()

    doc.add_heading(title, level=0)
    # Adicionar uma tabela ao documento
    table = doc.add_table(rows=1, cols=n)
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    doc.save(file)


def save_all(grammar, files):
    grammar_f = zip(grammar, files)
    for palavras, file in grammar_f:
        salva_palavras(palavras, file)

grammar = [verbs, adjectives, pronouns, noun, conjuction]
files = ['verbs.docx', 'adjectives.docx', 'pronouns.docx', 'noun.docx', 'conjuction.docx']

save_all(grammar, files)
