import nltk
import sys
import re
import os
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.cluster.util import cosine_distance
from nltk.corpus import stopwords
from youtube_transcript_api import YouTubeTranscriptApi
import numpy as np
import networkx as nx
from docx import Document
from docx.shared import Inches


# id do video
id = sys.argv[1]
# nome do arquivo
file = sys.argv[2]

srt = YouTubeTranscriptApi.get_transcript(str(id), languages=['en'])

texto = ''

for text in srt:
    texto = texto + text['text'] 


def remove_special_characters(text):
    # Substituir \xa0\n e \xa0\xa0 por espaços
    cleaned_text = re.sub(r'[\xa0\n]+', ' ', text)
    return cleaned_text


#removendo caracteris especiais
texto = remove_special_characters(texto)


def sentence_similarity(sent1, sent2, stopwords=None):
    if stopwords is None:
        stopwords = []
    
    words1 = [word.lower() for word in sent1 if word.isalnum() and word.lower() not in stopwords]
    words2 = [word.lower() for word in sent2 if word.isalnum() and word.lower() not in stopwords]
    
    all_words = list(set(words1 + words2))
    
    vector1 = [0] * len(all_words)
    vector2 = [0] * len(all_words)
    
    for word in words1:
        vector1[all_words.index(word)] += 1
        
    for word in words2:
        vector2[all_words.index(word)] += 1
        
    return 1 - cosine_distance(vector1, vector2)



def build_similarity_matrix(sentences, stop_words):
    similarity_matrix = np.zeros((len(sentences), len(sentences)))
    
    for idx1 in range(len(sentences)):
        for idx2 in range(len(sentences)):
            if idx1 == idx2:
                continue
            similarity_matrix[idx1][idx2] = sentence_similarity(sentences[idx1], sentences[idx2], stop_words)
    
    return similarity_matrix


def textrank(sentences, num_sentences=2):
    stop_words = set(stopwords.words('english'))
    sentence_similarity_matrix = build_similarity_matrix(sentences, stop_words)
    
    sentence_similarity_graph = nx.from_numpy_array(sentence_similarity_matrix)
    scores = nx.pagerank(sentence_similarity_graph)
    
    ranked_sentences = sorted(((scores[i], s) for i, s in enumerate(sentences)), reverse=True)
    
    top_sentences = []
    for i in range(num_sentences):
        top_sentences.append(ranked_sentences[i][1])
    
    return top_sentences


def salva_palavras(sentences, file):
    if os.path.isfile(file):
        doc = Document(file)
    else:
        doc = Document()

    for sentence in sentences:
        doc.add_paragraph(sentence)
    
    doc.save(file)

sentencas = sent_tokenize(texto)
palavras = [word_tokenize(sent) for sent in sentencas]

sentencas_importantes = textrank(sentencas, num_sentences=10)


salva_palavras(sentencas_importantes, file)
